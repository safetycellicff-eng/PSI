"""Builds the official PSI letter as a Word document.

Reproduces the Integral Coach Factory PSI report format: centred factory
heading, letter number, office/date/(SAFETY CELL) block, subject line, the
details table (Cause / Location / ShopNo / Description / Observation /
Action By), the violation photos, the Chief Safety Officer signature and the
"Copy to:-" table.

Edit the constants below if the office names or copy-to list ever change.
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

FACTORY_NAME = "INTEGRAL COACH FACTORY"
FACTORY_CITY = "CHENNAI 600038"
LETTER_NO = "NO: MFGL/Safety/PSI"
OFFICE_LINE = "OFFICE OF THE CWE/FUR"
CELL_LINE = "(SAFETY CELL)"
ADDRESSEE = "SSE/M/Fur"
SUBJECT = "Sub: -Plant safety inspection report in Furnishing Division -Reg."
SIGNATURE = "Chief Safety Officer/Fur"
COPY_TO = [
    ("CWE/LHB & Occupier", "For Kind Information please."),
    ("CSO/Fur ,Dy CEE/Maintenance", "For Kind Information please."),
]

CAUSE_NAMES = {
    "SV": "Safety Violation",
    "UA": "Unsafe Act",
    "UC": "Unsafe Condition",
    "NM": "Near Miss",
}


def build_psi_letter(record, photos):
    """Return .docx bytes for one PSI in the official letter format.

    record: dict keyed by storage.RECORD_HEADERS.
    photos: list of image bytes (the violation photos).
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Centred factory heading
    for text in (FACTORY_NAME, FACTORY_CITY):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(13)

    # Letter number
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(0)
    para.add_run(LETTER_NO)

    # Office / date / safety cell block, indented to the right half
    date_line = f"DATE: {datetime.now().strftime('%d-%m-%Y')}."
    for text in (OFFICE_LINE, date_line, CELL_LINE):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(3.6)
        para.paragraph_format.space_after = Pt(0)
        para.add_run(text)

    # Addressee
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(0)
    para.add_run(ADDRESSEE)

    # Subject
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)
    para.add_run(SUBJECT)

    _details_table(doc, record)
    _photos(doc, photos)

    # Signature, right aligned
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    para.add_run(SIGNATURE)

    _copy_to_table(doc)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _set_cell(cell, text, bold=False, runs=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if runs is None:
        runs = [(text, bold)]
    for run_text, run_bold in runs:
        run = para.add_run(run_text)
        run.bold = run_bold
        run.font.size = Pt(11.5)


def _details_table(doc, record):
    cause = CAUSE_NAMES.get(record.get("Category", ""), record.get("Category", ""))
    location = record.get("Location/Shop", "")
    date = record.get("First Appeared On", "")

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    widths = (Inches(2.35), Inches(2.35), Inches(1.95))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    _set_cell(table.cell(0, 0), "Cause", bold=True)
    _set_cell(table.cell(0, 1).merge(table.cell(0, 2)), cause)

    _set_cell(table.cell(1, 0), "Location", bold=True)
    _set_cell(table.cell(1, 1), location)
    _set_cell(table.cell(1, 2), None, runs=[("Date", True), (f" :- {date}", False)])

    _set_cell(table.cell(2, 0), "ShopNo / Firm Name", bold=True)
    _set_cell(table.cell(2, 1), location)
    _set_cell(table.cell(2, 2), None,
              runs=[("Shop Control : ", True), (record.get("Department", ""), False)])

    _set_cell(table.cell(3, 0), "Description of Violation / Hazard", bold=True)
    _set_cell(table.cell(3, 1).merge(table.cell(3, 2)),
              record.get("Description of Violation/Hazard", ""))

    _set_cell(table.cell(4, 0), "Observation / Suggestions", bold=True)
    _set_cell(table.cell(4, 1).merge(table.cell(4, 2)), record.get("Remarks", ""))

    _set_cell(table.cell(5, 0), "Action By", bold=True)
    _set_cell(table.cell(5, 1).merge(table.cell(5, 2)), record.get("Action By", ""))


def _photos(doc, photos, per_row=2, width=Inches(2.7)):
    if not photos:
        return
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    for start in range(0, min(len(photos), 4), per_row):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        for photo_bytes in photos[start : start + per_row]:
            run = para.add_run()
            run.add_picture(io.BytesIO(photo_bytes), width=width)
            para.add_run("  ")


def _copy_to_table(doc):
    table = doc.add_table(rows=1 + len(COPY_TO), cols=2)
    table.style = "Table Grid"
    header = table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell(header, "Copy to:-", bold=True)
    for i, (who, note) in enumerate(COPY_TO, start=1):
        table.cell(i, 0).width = Inches(2.9)
        table.cell(i, 1).width = Inches(3.75)
        _set_cell(table.cell(i, 0), who)
        _set_cell(table.cell(i, 1), note)
